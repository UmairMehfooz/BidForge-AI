"""
BidForge AI — Document Parser
==============================
Extracts plain text from RFP/RFQ/Tender documents (PDF or DOCX).

Public API
----------
    parse_document(file_path: str | Path) -> str

Design decisions
----------------
- PDF  : PyMuPDF (fitz) — page.get_text("text") only.  No layout mode, no OCR.
         Pages joined with \\n\\n.  If a page yields no text, it is silently skipped.
         If total extracted text < MIN_TEXT_CHARS, a 400 HTTPException is raised
         so the caller knows the PDF is likely scanned/image-only.
- DOCX : python-docx — iterates paragraphs, skips empty ones, joins with \\n\\n.
         Also extracts text from tables (cell by cell).
- Errors are surfaced as HTTPException with actionable detail messages so the
  FastAPI router can return them directly to the client.
"""

from __future__ import annotations

import logging
from pathlib import Path

from fastapi import HTTPException, status

logger = logging.getLogger(__name__)

# Minimum total characters for a PDF to be considered parseable.
# Anything less is almost certainly a scanned-image PDF.
MIN_TEXT_CHARS = 200

# Supported extensions → parser dispatch map
_SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_document(file_path: str | Path) -> str:
    """
    Extract and return the full plain text of a PDF or DOCX file.

    Parameters
    ----------
    file_path : str | Path
        Absolute or relative path to the RFP document.

    Returns
    -------
    str
        The full raw text of the document (pages / paragraphs joined with \\n\\n).

    Raises
    ------
    HTTPException 400
        - Unsupported file type.
        - PDF yielded fewer than MIN_TEXT_CHARS characters (likely scanned).
        - DOCX contained no readable paragraphs.
    HTTPException 422
        - File not found or unreadable at the given path.
    HTTPException 500
        - Unexpected parsing error (fitz crash, corrupt file, etc.).
    """
    path = Path(file_path)

    # ── existence guard ──────────────────────────────────────────────────────
    if not path.exists():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Document not found at path: {path}",
        )

    ext = path.suffix.lower()

    if ext not in _SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Unsupported file type '{ext}'. "
                f"Only {', '.join(sorted(_SUPPORTED_EXTENSIONS))} files are accepted."
            ),
        )

    if ext == ".pdf":
        return _parse_pdf(path)
    else:
        return _parse_docx(path)


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> str:
    """
    Extract text from a PDF using pypdf.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="pypdf is not installed. Run: pip install pypdf",
        )

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        logger.error("PdfReader failed for %s: %s", path, exc)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Could not open PDF '{path.name}': {exc}",
        )

    page_texts: list[str] = []
    empty_pages: list[int] = []

    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            # Page markers let the requirement extractor map chunks back to
            # page ranges (Fix 5). Stripped before any text reaches the LLM.
            page_texts.append(f"[[PAGE {page_num + 1}]]\n{text.strip()}")
        else:
            empty_pages.append(page_num + 1)
            logger.warning(
                "PDF '%s' page %d yielded no text (possibly image-only).",
                path.name, page_num + 1,
            )

    full_text = "\n\n".join(page_texts)

    # ── Minimum content guard ────────────────────────────────────────────────
    if len(full_text.strip()) < MIN_TEXT_CHARS:
        empty_summary = f" (empty pages: {empty_pages})" if empty_pages else ""
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"PDF '{path.name}' yielded only {len(full_text.strip())} characters"
                f"{empty_summary}. "
                "The file appears to be a scanned image PDF. "
                "Please provide a text-selectable PDF."
            ),
        )

    logger.info(
        "PDF parsed: '%s' — %d pages, %d chars extracted (%d empty pages skipped).",
        path.name, len(page_texts) + len(empty_pages),
        len(full_text), len(empty_pages),
    )
    return full_text



# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------

def _parse_docx(path: Path) -> str:
    """
    Extract text from a DOCX using python-docx.

    Extracts:
    - All non-empty paragraphs (body text, headings, lists).
    - All table cell text (row by row, cell by cell).

    Returns paragraphs + table text joined by \\n\\n.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="python-docx is not installed. Run: pip install python-docx",
        )

    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.error("Document() failed for %s: %s", path, exc)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Could not open DOCX '{path.name}': {exc}",
        )

    blocks: list[str] = []

    # ── Body paragraphs ──────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    # ── Tables ───────────────────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                blocks.append(" | ".join(row_cells))

    if not blocks:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"DOCX '{path.name}' contained no readable text. "
                "Please verify the file is not empty or protected."
            ),
        )

    full_text = "\n\n".join(blocks)
    logger.info(
        "DOCX parsed: '%s' — %d blocks, %d chars extracted.",
        path.name, len(blocks), len(full_text),
    )
    return full_text


# ---------------------------------------------------------------------------
# Smoke test — run directly:  python -m app.services.document_parser <file>
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s  %(message)s")

    if len(sys.argv) < 2:
        print("Usage: python -m app.services.document_parser <path/to/file.pdf|docx>")
        sys.exit(1)

    target = sys.argv[1]
    print(f"\nParsing: {target}\n{'='*60}")
    try:
        text = parse_document(target)
        preview = text[:1000]
        print(preview)
        print(f"\n{'='*60}")
        print(f"Total characters extracted: {len(text)}")
    except HTTPException as e:
        print(f"HTTPException {e.status_code}: {e.detail}")
        sys.exit(1)
