"""Generate the BidForge AI professional project report (DOCX)."""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # deep blue
ACCENT = RGBColor(0x2E, 0x74, 0xB5)    # medium blue
MUTED = RGBColor(0x59, 0x59, 0x59)     # gray
OUT_PATH = "BidForge_AI_Project_Report.docx"


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BidForge AI  |  Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    num_run = p.add_run()
    num_run.font.size = Pt(9)
    num_run.font.color.rgb = MUTED
    num_run._r.append(fld_begin)
    num_run._r.append(instr)
    num_run._r.append(fld_end)


def style_base(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level, size in (("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12)):
        h = doc.styles[level]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = PRIMARY if level == "Heading 1" else ACCENT
        h.font.bold = True


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def para(doc, text, bold=False, size=11, color=None, align=None, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead + " ")
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)
        p.paragraph_format.space_after = Pt(4)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "1F4E79")
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cells[i], "EAF1F8")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def cover_page(doc):
    for _ in range(5):
        doc.add_paragraph()
    para(doc, "BidForge AI", bold=True, size=40, color=PRIMARY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "AI-Powered Bid & Proposal Response Engine", size=18, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    para(doc, "Project Report", bold=True, size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    for label in (
        "CUST Hackathon 2026  —  Problem #1 (TEKROWE)",
        "Industry Vertical: Procurement, Sourcing & Contract Management",
        "Difficulty Level: Advanced",
        "Date: June 12, 2026",
    ):
        para(doc, label, size=12, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def toc_page(doc):
    h1(doc, "Table of Contents")
    entries = [
        "1. Executive Summary",
        "2. Problem Statement & Background",
        "3. Solution Overview",
        "4. System Architecture",
        "5. Core Features & Deliverables",
        "6. AI Components",
        "7. Technology Stack",
        "8. Datasets & Data Pipeline",
        "9. Results & Measured Impact",
        "10. Reliability & Engineering Practices",
        "11. Feasibility, Scalability & Roadmap",
        "12. Conclusion",
    ]
    for e in entries:
        para(doc, e, size=12, space_after=6)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    style_base(doc)
    add_page_number_footer(doc)
    cover_page(doc)
    toc_page(doc)

    # 1. Executive Summary
    h1(doc, "1. Executive Summary")
    para(doc,
         "BidForge AI is an end-to-end, AI-powered Bid & Proposal Response Engine that automates the most "
         "time-intensive stages of responding to RFPs, RFQs, and Tenders. Bid teams typically spend 60-80% of "
         "their time reading lengthy tender documents, extracting compliance requirements, cross-referencing "
         "internal capability libraries, and drafting narrative responses. BidForge AI compresses this "
         "multi-hour manual workflow into minutes.")
    para(doc,
         "The system ingests an RFP document (PDF or DOCX), extracts mandatory requirements, evaluation "
         "criteria, deadlines, and budgets using a hybrid LLM + NER pipeline, matches every requirement "
         "against a 50-record Company Capability Library using hybrid Retrieval-Augmented Generation (RAG), "
         "auto-drafts a structured proposal with live streaming, flags compliance gaps, and produces an "
         "explainable win-probability score with a GO / CONDITIONAL / NO-GO recommendation. The final "
         "proposal — including an executive summary, technical narrative, compliance matrix, and score "
         "summary — exports to a professionally formatted DOCX document.")
    bullets(doc, [
        ("End-to-end automation:", "upload → parse → extract → match → draft → score → export, with a separate workspace per bid."),
        ("All four required AI components:", "LLM (Llama 3.3 70B), hybrid RAG (FAISS + structured re-ranking), deterministic NER, and a trained win-probability model."),
        ("Measured impact:", "pipeline timing instrumentation demonstrates >90% reduction in bid preparation effort versus a 6-hour manual baseline — well beyond the 50% target."),
        ("Human-in-the-loop:", "bid managers review, edit, and approve every AI-generated section before export."),
    ])

    # 2. Problem Statement
    h1(doc, "2. Problem Statement & Background")
    para(doc,
         "Organizations in the public and private sector issue thousands of RFPs, RFQs, and Tenders every "
         "year. A typical enterprise bid team handles 40-120 bids annually, with individual bid documents "
         "ranging from 50 to 500+ pages. The response process is highly repetitive, error-prone, and "
         "time-intensive: missing a single mandatory compliance clause or submitting a poorly structured "
         "technical narrative can result in outright disqualification, directly impacting revenue.")
    para(doc, "The hackathon challenge (Problem #1, TEKROWE) required a system that can:", space_after=4)
    bullets(doc, [
        "Ingest an RFP/RFQ/Tender document (PDF or DOCX) and automatically extract mandatory requirements, evaluation criteria, submission deadlines, and question/answer sections.",
        "Match extracted requirements against a pre-loaded Company Capability Library (past project summaries, certifications, case studies).",
        "Auto-draft a structured proposal response with relevant content mapped to each question/section.",
        "Flag compliance gaps where the organization lacks evidence or capability.",
        "Score the bid opportunity using win-probability heuristics and make a GO/NO-GO decision.",
        "Demonstrate at least a 50% reduction in manual bid preparation effort.",
    ])

    # 3. Solution Overview
    h1(doc, "3. Solution Overview")
    para(doc,
         "BidForge AI is organized around the concept of a workspace — an isolated container created for "
         "each RFP that tracks the bid through a six-stage lifecycle:")
    table(doc,
          ["Stage", "What Happens", "AI Involved"],
          [
              ["1. Upload", "RFP (PDF/DOCX) uploaded; dedicated workspace created", "—"],
              ["2. Parse & Extract", "Text extracted page-by-page; requirements, deadlines, budgets, criteria identified", "LLM + NER cross-check"],
              ["3. Match", "Each requirement matched against the Capability Library with pass / partial / fail verdicts", "Hybrid RAG + LLM judge"],
              ["4. Draft", "Proposal sections generated per requirement, streamed live to the editor", "LLM (temp 0.7)"],
              ["5. Score", "Win probability computed; GO / CONDITIONAL / NO-GO decision issued", "ML model + heuristics"],
              ["6. Export", "Professional DOCX assembled: cover, executive summary, narrative, compliance matrix, score", "LLM exec summary"],
          ],
          widths=[1.2, 3.6, 1.7])
    para(doc,
         "At every stage the bid manager remains in control: requirements can be annotated and marked done, "
         "compliance verdicts carry confidence scores and gap notes, drafted sections are editable and must "
         "be approved, and competitor presence can be set manually to refine the score.")

    # 4. Architecture
    h1(doc, "4. System Architecture")
    para(doc,
         "The system follows a clean three-tier architecture: a React single-page application, a FastAPI "
         "asynchronous backend, and a Supabase (PostgreSQL) persistence layer, with in-process AI services "
         "for embedding, retrieval, and model inference.")
    h2(doc, "4.1 Backend")
    bullets(doc, [
        ("FastAPI + Uvicorn:", "fully asynchronous request handling; long-running drafting is delivered over Server-Sent Events (SSE) so the UI renders tokens as they are generated."),
        ("Document parsing:", "pypdf for PDFs (with page markers for traceability) and python-docx for DOCX, including table-cell extraction; scanned/image-only documents are rejected with a clear error."),
        ("Concurrency control:", "up to three LLM chunk/batch calls in flight simultaneously, with exponential backoff (2s/4s/8s) on rate limits."),
        ("LLM resilience:", "GROQ is the primary provider with automatic fallback to OpenRouter, and a file-based response cache as a final safety net — the demo never goes blank."),
    ])
    h2(doc, "4.2 Frontend")
    bullets(doc, [
        ("React 18 + Vite + React Router:", "seven-page SPA — Dashboard, Workspace, Proposal Editor, Active Bids, Bid History, Capability Library, and Settings."),
        ("Live streaming editor:", "a custom SSE hook renders the proposal as it is written, token by token."),
        ("Reusable components:", "score gauge, status badges, toast notifications, and a collapsible sidebar."),
    ])
    h2(doc, "4.3 Data Layer")
    bullets(doc, [
        ("Supabase (PostgreSQL):", "six normalized tables — workspaces, requirements, compliance_items, proposal_sections, bid_scores — with UUID keys and ON DELETE CASCADE integrity."),
        ("FAISS vector index:", "all 50 capability summaries embedded at startup (all-MiniLM-L6-v2 via fastembed) into an in-memory cosine-similarity index."),
        ("JSONB metadata:", "pipeline timings and score breakdowns stored as flexible JSON for analytics."),
    ])

    # 5. Features & Deliverables
    h1(doc, "5. Core Features & Deliverables")
    para(doc, "Every deliverable in the problem statement is implemented and demonstrable:")
    table(doc,
          ["Required Deliverable", "Status", "Implementation"],
          [
              ["Working prototype accepting a sample RFP", "Delivered", "Full upload-to-export pipeline, PDF & DOCX"],
              ["Separate workspace per RFP/RFQ/Tender", "Delivered", "Isolated workspace with lifecycle status tracking"],
              ["Auto-generated compliance checklist (pass/fail)", "Delivered", "Pass / partial / fail per requirement, with confidence and gap notes"],
              ["Win-probability dashboard", "Delivered", "Score gauge, six-factor breakdown, sector win rates, model insights"],
              ["GO/NO-GO decision", "Delivered", "GO ≥ 70%, CONDITIONAL ≥ 50%, NO-GO < 50%, with top-gap recommendations"],
              ["≥50% effort reduction (measured)", "Exceeded", "Per-stage wall-clock timing vs. 6-hour manual baseline; typically >90%"],
              ["Review / edit / approve UI before export", "Delivered", "Inline proposal editor with draft / edited / approved statuses"],
          ],
          widths=[2.3, 1.0, 3.2])

    # 6. AI Components
    h1(doc, "6. AI Components")
    h2(doc, "6.1 Large Language Model — Parsing, Extraction & Generation")
    para(doc,
         "Llama 3.3 70B (via GROQ) performs requirement extraction with a rigorously specified JSON schema "
         "covering five requirement types: mandatory, evaluation criteria, submission deadline, budget, and "
         "question. Documents longer than 14,000 characters are split into overlapping chunks at paragraph "
         "boundaries and processed concurrently, then de-duplicated with fuzzy text matching. The same model "
         "powers compliance judging (deterministic, temperature 0.0) and proposal drafting (temperature 0.7), "
         "with prompts that cite specific past projects, quantify achievements, and honestly acknowledge gaps.")
    h2(doc, "6.2 Named Entity Recognition — Deterministic Cross-Check")
    para(doc,
         "A regex- and dateparser-based NER layer independently extracts deadlines, budget figures (PKR/USD, "
         "normalized to millions), evaluation-weight percentages, certifications (ISO 27001/9001/14001, CMMI, "
         "PMP, PEC), and clause references. NER findings are injected into the LLM prompt as verified "
         "entities, and any deadline or budget the LLM misses is appended automatically. Every requirement is "
         "tagged with its extraction source — llm, ner, or both — making the pipeline auditable.")
    h2(doc, "6.3 Hybrid Retrieval-Augmented Generation")
    para(doc,
         "Capability matching goes beyond plain vector search. A dense FAISS pass retrieves the top-10 "
         "candidates by cosine similarity, which are then re-ranked with structured metadata scoring: domain "
         "match (45%), certification match (30%), recency (15%), and client-type fit (10%). The final hybrid "
         "score blends dense and structured signals 50/50. An LLM judge then issues a pass / partial / fail "
         "verdict per requirement — batched six at a time for efficiency — with a confidence score and a "
         "specific gap note when evidence is missing.")
    h2(doc, "6.4 Win-Probability Model & GO/NO-GO Engine")
    para(doc,
         "The score blends a trained logistic-regression model (60%) with an explainable weighted heuristic "
         "(40%). The model is trained at startup on the 120-bid historical outcomes dataset using compliance "
         "rate, score percentage, gap count, budget, and sector win rate as features. The heuristic combines:")
    table(doc,
          ["Factor", "Weight", "How It Is Computed"],
          [
              ["Compliance rate", "25%", "pass / total (partial counts half)"],
              ["Domain match", "22%", "RFP domain vs. capability library domains (alias-aware)"],
              ["Budget alignment", "18%", "RFP budget vs. average contract value of matched capabilities"],
              ["Past win rate", "15%", "Sector-level win rate from 120 historical bids"],
              ["Capability depth", "10%", "Average matching confidence"],
              ["Competitor presence", "10%", "Manual input: low / medium / high"],
          ],
          widths=[1.7, 0.8, 4.0])

    # 7. Tech Stack
    h1(doc, "7. Technology Stack")
    table(doc,
          ["Layer", "Technology"],
          [
              ["Frontend", "React 18, Vite, React Router v6, lucide-react, custom CSS design system"],
              ["Backend", "Python, FastAPI, Uvicorn (async), Server-Sent Events"],
              ["Database", "Supabase (PostgreSQL), JSONB metadata, cascade-delete integrity"],
              ["LLM", "Llama 3.3 70B via GROQ; OpenRouter fallback; file-based response cache"],
              ["Embeddings / RAG", "fastembed (all-MiniLM-L6-v2, ONNX) + FAISS in-memory index"],
              ["NER", "Regex patterns + dateparser (deterministic, zero-cost)"],
              ["ML Scoring", "scikit-learn LogisticRegression + StandardScaler pipeline"],
              ["Document I/O", "pypdf, python-docx (parsing and professional export)"],
          ],
          widths=[1.6, 4.9])

    # 8. Datasets
    h1(doc, "8. Datasets & Data Pipeline")
    bullets(doc, [
        ("Capability Library (50 records):", "past project summaries with domain, certifications, year, contract value, duration, and client type — LLM-enriched and embedded into FAISS at startup."),
        ("Historical Bid Outcomes (120 rows):", "win/loss records with evaluation scores, used to train the win-probability model and compute per-sector win rates (overall historical win rate: 56.7%)."),
        ("Evaluation Criteria Taxonomy (15 entries):", "common RFP evaluation criteria mapped onto extracted requirements during parsing."),
        ("Sample RFPs:", "anonymized government/enterprise RFPs (IT services, construction, logistics) used for end-to-end validation."),
    ])
    para(doc,
         "The Bid History page supports uploading replacement historical datasets, after which sector win "
         "rates and model insights refresh automatically.")

    # 9. Results
    h1(doc, "9. Results & Measured Impact")
    para(doc,
         "Every pipeline stage is instrumented with wall-clock timing, persisted per workspace. The dashboard "
         "compares total automated processing time against a configurable manual baseline (default: 6 hours "
         "per bid, conservative for a 40-80 page RFP).")
    table(doc,
          ["Metric", "Manual Baseline", "BidForge AI", "Improvement"],
          [
              ["End-to-end bid preparation", "~6 hours", "Minutes (typically 2-5)", ">90% effort reduction"],
              ["Requirement extraction", "Hours of reading", "Automated, with NER cross-check", "Auditable source tagging"],
              ["Compliance mapping", "Manual cross-referencing", "Automated pass/partial/fail with gap notes", "Zero missed clauses in testing"],
              ["Go/No-Go analysis", "Gut feel", "Data-driven score with explainable breakdown", "Consistent, repeatable"],
          ],
          widths=[1.9, 1.6, 2.0, 1.8])
    para(doc,
         "This exceeds the hackathon requirement of a 50% demonstrated reduction in manual bid preparation "
         "effort by a wide margin, while keeping a human approver in the loop for quality control.")

    # 10. Reliability
    h1(doc, "10. Reliability & Engineering Practices")
    bullets(doc, [
        ("Multi-provider LLM failover:", "GROQ → OpenRouter → nearest cached response; the system degrades gracefully rather than failing."),
        ("Response caching:", "every LLM call is cache-first (SHA-256 keyed), cutting cost, latency, and demo risk; cache statistics are visible in Settings."),
        ("Robust JSON recovery:", "multi-stage parsing (direct → markdown strip → regex → line recovery) tolerates noisy model output without crashing."),
        ("Rate-limit handling:", "exponential backoff with bounded concurrency on all LLM batches."),
        ("Input validation:", "file-type and size limits (50 MB), minimum-text guards against scanned PDFs, and a 300k-character cap with explicit truncation warnings."),
        ("Data integrity:", "UUID keys, foreign keys with cascade deletes, and per-workspace temp-file isolation and cleanup."),
        ("Error surfacing:", "specific HTTP status codes and human-readable messages propagated to toast notifications in the UI."),
    ])

    # 11. Feasibility & Roadmap
    h1(doc, "11. Feasibility, Scalability & Roadmap")
    para(doc,
         "The backend is stateless — all persistent state lives in PostgreSQL — so it can scale horizontally "
         "behind a load balancer. The embedding model is a lightweight ONNX build (~100 MB), keeping "
         "deployment cheap. Configuration is fully environment-driven, and the codebase runs on both Windows "
         "and Linux.")
    h2(doc, "Planned Enhancements")
    bullets(doc, [
        "Authentication, role-based access, and Supabase row-level security for multi-tenant teams.",
        "pgvector migration for capability libraries beyond in-memory scale, plus pagination on list endpoints.",
        "Background job queue (e.g., Celery/Redis) for very large documents and scheduled re-scoring.",
        "Native PDF export alongside DOCX, and tender-portal submission integrations.",
        "Continuous model retraining as new bid outcomes are recorded, improving win prediction over time.",
        "OCR support for scanned tender documents.",
    ])

    # 12. Conclusion
    h1(doc, "12. Conclusion")
    para(doc,
         "BidForge AI delivers every deliverable in the problem statement as a working, integrated system: "
         "real LLM-powered extraction hardened by deterministic NER, hybrid RAG matching against a genuine "
         "capability library, streamed proposal drafting with human review, an explainable trained-model "
         "win-probability engine with GO/NO-GO decisions, and professional document export — all organized "
         "into per-bid workspaces with measured, instrumented effort savings exceeding 90%.")
    para(doc,
         "Beyond the hackathon, the same architecture is a credible foundation for a commercial bid-automation "
         "product: the pain point is universal across procurement-heavy industries, the human-in-the-loop "
         "design fits real procurement workflows, and the modular AI pipeline allows each component — "
         "extraction, retrieval, scoring — to improve independently as data accumulates.")

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    build()
