"""
Generate BidForge_AI_Documentation.docx — full project documentation.
Run:  .\\venv\\Scripts\\python.exe scripts\\generate_documentation.py
"""

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).parent.parent
OUT = ROOT / "BidForge_AI_Documentation.docx"

NAVY = RGBColor(0, 51, 102)
GREY = RGBColor(110, 110, 110)


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph("BidForge AI")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.font.size = Pt(40)
        r.bold = True
        r.font.color.rgb = NAVY
    s = doc.add_paragraph("AI-Powered RFP Analysis & Proposal Generation Platform")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in s.runs:
        r.font.size = Pt(16)
    doc.add_paragraph()
    v = doc.add_paragraph("Technical & User Documentation")
    v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in v.runs:
        r.font.size = Pt(13)
        r.font.italic = True
    for _ in range(6):
        doc.add_paragraph()
    d = doc.add_paragraph(f"Version 1.0  ·  {datetime.now().strftime('%d %B %Y')}")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in d.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = GREY
    doc.add_page_break()


def toc(doc):
    doc.add_heading("Contents", level=1)
    items = [
        "1. Overview",
        "2. System Architecture",
        "3. Getting Started — Installation & Setup",
        "4. The Bid Pipeline (How a Bid Flows Through the System)",
        "5. Application Pages (User Guide)",
        "6. Document Exports",
        "7. API Reference",
        "8. Capability Library",
        "9. LLM Strategy — Caching, Fallback & Demo Mode",
        "10. Database Schema & Migrations",
        "11. Troubleshooting & Known Behaviours",
    ]
    for i in items:
        doc.add_paragraph(i, style="List Number" if False else None).paragraph_format.left_indent = Pt(12)
    doc.add_page_break()


def para(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc, headers, rows, style="Light Shading Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build():
    doc = Document()
    cover(doc)
    toc(doc)

    # ── 1. Overview ─────────────────────────────────────────────────────────
    doc.add_heading("1. Overview", level=1)
    para(doc, (
        "BidForge AI is an end-to-end bid management platform for companies that respond "
        "to RFPs (Requests for Proposal). It ingests an RFP document, extracts every "
        "requirement automatically, checks each one against the company's documented "
        "capability library, drafts proposal narrative with AI, computes a win-probability "
        "score with a GO/NO-GO recommendation, and exports a formatted proposal as a "
        "Microsoft Word document."
    ))
    para(doc, "The platform answers four questions for every incoming RFP:")
    bullets(doc, [
        "What exactly is the client asking for? (Parse — requirement extraction)",
        "Can we prove we have done this before? (Match — RAG compliance matching)",
        "What should our response say? (Draft — AI proposal generation with evidence citations)",
        "Should we bid at all? (Score — win-probability model trained on historical bid outcomes)",
    ])
    doc.add_page_break()

    # ── 2. Architecture ─────────────────────────────────────────────────────
    doc.add_heading("2. System Architecture", level=1)
    table(doc, ["Layer", "Technology", "Role"], [
        ["Frontend", "React 18 + Vite + Tailwind CSS", "Single-page app: dashboard, workspace, proposal editor, library management"],
        ["Backend", "FastAPI (Python 3.12) + Uvicorn", "REST API + Server-Sent Events streaming, all pipeline logic"],
        ["Database", "Supabase (PostgreSQL)", "Workspaces, requirements, compliance items, proposal sections, bid scores"],
        ["Vector search", "FAISS + fastembed (all-MiniLM-L6-v2, ONNX)", "Dense retrieval over the capability library; hybrid re-ranking with metadata"],
        ["LLM (primary)", "GROQ — llama-3.3-70b-versatile", "Requirement extraction, compliance matching, proposal drafting, summaries"],
        ["LLM (fallback)", "OpenRouter free-model rotation", "Automatic failover when GROQ is rate-limited or down"],
        ["Win model", "scikit-learn LogisticRegression", "Win probability from 6 factors, trained on data/bid_history.csv"],
        ["Document output", "python-docx", "Draft export and full 8-section proposal export"],
    ])
    doc.add_paragraph()
    doc.add_heading("Project layout", level=2)
    code(doc, (
        "BidForge AI/\n"
        "  app/\n"
        "    main.py                  FastAPI app, startup, /health\n"
        "    routers/                 workspaces, bid_history, capabilities, settings\n"
        "    services/                document_parser, requirement_extractor,\n"
        "                             compliance_matcher, proposal_generator,\n"
        "                             scoring_engine, win_model, docx_exporter,\n"
        "                             capability_store, llm_fallback ...\n"
        "    models/schemas.py        Pydantic request/response models\n"
        "    utils/                   supabase_client, llm_cache\n"
        "    data/                    capability_library(.enriched).json, cache/\n"
        "  frontend/                  React app (Vite)\n"
        "  scripts/                   warm_cache, enrich_capabilities, utilities\n"
        "  schema.sql                 Supabase DDL + migration notes\n"
        "  venv/                      Python virtual environment\n"
        "  .env                       API keys & runtime flags"
    ))
    doc.add_page_break()

    # ── 3. Setup ────────────────────────────────────────────────────────────
    doc.add_heading("3. Getting Started — Installation & Setup", level=1)
    doc.add_heading("3.1 Prerequisites", level=2)
    bullets(doc, [
        "Python 3.12 (the project venv lives at .\\venv — there is no system Python on the dev machine)",
        "Node.js 18+ (frontend)",
        "A Supabase project (run schema.sql in its SQL editor)",
        "A GROQ API key (free tier works; 100k tokens/day cap)",
        "Optionally an OpenRouter API key for automatic failover",
    ])
    doc.add_heading("3.2 Environment variables (.env)", level=2)
    table(doc, ["Variable", "Purpose"], [
        ["SUPABASE_URL / SUPABASE_KEY", "Supabase project connection"],
        ["GROQ_API_KEY", "Primary LLM provider"],
        ["OPENROUTER_API_KEY", "Fallback LLM rotation (optional but recommended)"],
        ["DEMO_MODE", "true = serve nearest cached LLM response when all providers fail (judge-day insurance)"],
        ["MANUAL_BASELINE_HOURS", "Manual-effort baseline used for the time-saved metric"],
    ])
    doc.add_heading("3.3 Running the backend", level=2)
    code(doc, ".\\venv\\Scripts\\python.exe -m uvicorn app.main:app --host 0.0.0.0 --port 8000")
    para(doc, "Interactive API docs are served at http://localhost:8000/docs; liveness at /health.")
    doc.add_heading("3.4 Running the frontend", level=2)
    code(doc, "cd frontend\nnpm install\nnpm run dev      # development, proxies /api to :8000\nnpm run build    # production bundle in frontend/dist")
    doc.add_page_break()

    # ── 4. Pipeline ─────────────────────────────────────────────────────────
    doc.add_heading("4. The Bid Pipeline", level=1)
    para(doc, (
        "Every bid lives in a workspace and moves through six stages. Each stage is one "
        "API call, triggered from the UI; the workspace status field tracks progress "
        "(created → uploaded → parsed → matched → drafted → scored → exported)."
    ))
    table(doc, ["Stage", "Endpoint", "What happens"], [
        ["1. Upload", "POST /api/workspaces", "PDF/DOCX RFP is uploaded and stored; a workspace row is created."],
        ["2. Parse", "POST /api/workspaces/{id}/parse", "Document text is chunked and sent to the LLM (plus an NER pass); structured requirements are extracted with section refs, types (mandatory / question / evaluation criteria / submission deadline / budget), deadlines, and budget references. Sector and issuer type are inferred."],
        ["3. Match", "POST /api/workspaces/{id}/match", "Each requirement is matched against the capability library with hybrid retrieval (50% dense cosine + 50% structured metadata: domain, certification, recency, client type). Produces PASS / PARTIAL / FAIL compliance items with confidence and gap notes."],
        ["4. Draft", "POST /api/workspaces/{id}/draft", "For every mandatory/question requirement, the top-3 matching capabilities are fed to the LLM, which writes a 2-3 paragraph evidence-grounded response. Sections stream token-by-token to the editor (SSE) and are batch-saved to the database."],
        ["5. Score", "POST /api/workspaces/{id}/score", "Six factors (compliance rate, domain match, budget alignment, past win rate, capability depth, competitor presence) feed a logistic-regression win model → overall score and GO / CONDITIONAL / NO-GO decision."],
        ["6. Export", "POST /api/workspaces/{id}/export or /export-full", "Generates the proposal DOCX (see section 6)."],
    ])
    doc.add_page_break()

    # ── 5. Pages ────────────────────────────────────────────────────────────
    doc.add_heading("5. Application Pages (User Guide)", level=1)

    doc.add_heading("5.1 Dashboard", level=2)
    bullets(doc, [
        "Stat cards: total bids, GO decisions, average win score, proposals exported, average time saved vs manual.",
        "Status filter chips: All / Pending / In Progress / Done.",
        "One card per bid showing sector, status, compliance bar, win score, and decision badge.",
        "Analyze New RFP opens the upload modal.",
        "Delete: hover a card and click the trash icon to permanently delete the bid and all of its data (a confirmation prompt is shown first).",
    ])

    doc.add_heading("5.2 Active Bids", level=2)
    bullets(doc, [
        "Table view of every workspace with search across name, sector, status, and decision.",
        "Click a row to open the workspace; the trash icon on each row deletes the bid.",
    ])

    doc.add_heading("5.3 Workspace (master-detail)", level=2)
    bullets(doc, [
        "Left column: every extracted requirement with section ref, type badge, and filter tabs (All / Mandatory / Evaluation / Deadline).",
        "Centre column: the selected requirement, its compliance analysis (status, matched capability, confidence bar, gap note), and the drafted proposal content for it.",
        "Right column: win-score gauge, GO/NO-GO badge, score breakdown bars, competitor-presence input (re-scores live), and top gaps.",
        "My Note: each requirement has a short note box — type a note and press Save note. Notes are saved to the database and shown as a small preview in the requirement list.",
        "Mark as done: a per-requirement toggle. Done requirements show a green check and strikethrough in the list, so you can track review progress through a long RFP.",
    ])

    doc.add_heading("5.4 Proposal Editor", level=2)
    bullets(doc, [
        "Generate Proposal Draft streams one AI-written section per requirement, live, token by token.",
        "Each section can be edited (Edit → Save Draft) and approved (Approve). Edits and approvals are saved to the database immediately; the approved state persists across refreshes.",
        "If you approve while a freshly generated draft is still being saved, the editor tells you to retry in a few seconds instead of faking success.",
        "Two export buttons (see section 6): Export Full Proposal and Export Draft DOCX. Buttons show progress and real error messages; exports can be repeated any number of times.",
    ])

    doc.add_heading("5.5 Bid History", level=2)
    bullets(doc, [
        "Shows the historical bid outcomes dataset (the win model's training data).",
        "Upload a replacement CSV to retrain the win model.",
    ])

    doc.add_heading("5.6 Capability Library", level=2)
    bullets(doc, [
        "Browse all capability records (past projects) the matcher uses as evidence.",
        "Upload a replacement JSON library; the FAISS index is rebuilt automatically.",
    ])

    doc.add_heading("5.7 Settings", level=2)
    bullets(doc, [
        "System status (providers, cache statistics) and runtime toggles persisted to .env.",
        "Clear LLM cache button.",
    ])
    doc.add_page_break()

    # ── 6. Exports ──────────────────────────────────────────────────────────
    doc.add_heading("6. Document Exports", level=1)
    para(doc, "Two DOCX exports are available from the Proposal Editor:")
    doc.add_heading("6.1 Export Draft DOCX  (POST /api/workspaces/{id}/export)", level=2)
    para(doc, (
        "A working draft: cover page, AI executive summary, the proposal narrative "
        "(every section in natural order, edited text preferred over raw AI draft), a "
        "compliance checklist table, and the win-score analysis page."
    ))
    doc.add_heading("6.2 Export Full Proposal  (POST /api/workspaces/{id}/export-full)", level=2)
    para(doc, "A complete, submission-shaped proposal with eight parts:")
    table(doc, ["#", "Section", "Source"], [
        ["1", "Cover Page", "Company name, RFP reference (workspace name), date"],
        ["2", "Executive Summary", "LLM-generated from the drafted sections"],
        ["3", "Understanding of Requirements", "LLM-written from the extracted requirements; deterministic fallback if providers fail"],
        ["4", "Technical Proposal", "One subsection per requirement: requirement text, your approved/edited draft (preferred over the raw AI draft), and a supporting-evidence citation. Requirements with no matching capability are flagged: [GAP FLAGGED - No evidence found for this requirement]"],
        ["5", "Company Profile & Credentials", "Aggregated deterministically from the capability library (practice areas, certifications, project counts) — no LLM, nothing can be fabricated"],
        ["6", "Relevant Past Projects", "Three case studies of the capabilities most frequently matched against this RFP, written strictly from library facts"],
        ["7", "Compliance Matrix", "Table: Requirement | Compliant (Yes/Partial/No) | Evidence"],
        ["8", "Appendices", "Placeholder list for registration documents, certificates, CVs, financials"],
    ])
    para(doc, (
        "Every export is written to a unique timestamped file, so repeated exports never "
        "collide with files that are still open. All LLM-written parts go through the "
        "cache and provider rotation — the first full export of a workspace takes about "
        "1-2 minutes; repeats take seconds."
    ))
    doc.add_page_break()

    # ── 7. API Reference ────────────────────────────────────────────────────
    doc.add_heading("7. API Reference", level=1)
    para(doc, "Base URL: http://localhost:8000 — interactive docs at /docs.")
    table(doc, ["Method", "Path", "Summary"], [
        ["GET", "/health", "Liveness probe: service status, capability count, win-model state, demo mode"],
        ["POST", "/api/workspaces", "Create a workspace and upload the RFP document"],
        ["GET", "/api/workspaces", "List all workspaces"],
        ["GET", "/api/workspaces/dashboard", "Bulk dashboard summary for all workspaces"],
        ["GET", "/api/workspaces/{id}", "Get one workspace"],
        ["PATCH", "/api/workspaces/{id}", "Update manual inputs (competitor presence)"],
        ["DELETE", "/api/workspaces/{id}", "Delete a bid and all its data (requirements, compliance, sections, scores, files)"],
        ["GET", "/api/workspaces/{id}/overview", "Workspace + requirements + compliance + sections + latest score"],
        ["POST", "/api/workspaces/{id}/parse", "Extract structured requirements from the RFP"],
        ["POST", "/api/workspaces/{id}/match", "RAG compliance matching against the capability library"],
        ["POST", "/api/workspaces/{id}/draft", "Stream AI proposal sections (Server-Sent Events)"],
        ["POST", "/api/workspaces/{id}/score", "Win-probability score and GO/NO-GO decision"],
        ["PATCH", "/api/workspaces/{id}/proposal/{section_id}", "Save human edits / approval on a proposal section"],
        ["PATCH", "/api/workspaces/{id}/requirements/{req_id}", "Save a note and/or done-mark on a requirement"],
        ["POST", "/api/workspaces/{id}/export", "Generate and download the draft proposal DOCX"],
        ["POST", "/api/workspaces/{id}/export-full", "Generate and download the complete 8-section proposal DOCX"],
        ["GET", "/api/bid-history", "Get the historical bid outcomes dataset"],
        ["POST", "/api/bid-history/upload", "Replace bid history CSV and retrain the win model"],
        ["GET", "/api/capabilities", "Get the capability library"],
        ["POST", "/api/capabilities/upload", "Replace the capability library and rebuild the index"],
        ["GET", "/api/settings", "System status and configuration"],
        ["PATCH", "/api/settings", "Update runtime configuration (persisted to .env)"],
        ["POST", "/api/settings/clear-cache", "Delete all cached LLM responses"],
        ["GET", "/api/debug/match?q=...", "Ad-hoc hybrid match against the library"],
        ["GET", "/api/debug/cache-stats", "LLM cache hit/miss/fallback statistics"],
    ])
    doc.add_page_break()

    # ── 8. Capability Library ───────────────────────────────────────────────
    doc.add_heading("8. Capability Library", level=1)
    para(doc, (
        "The library is the single source of evidence for compliance matching and "
        "proposal drafting: 60 past-project records in app/data/capability_library.json "
        "(an enriched variant with richer summaries is preferred when present). Each "
        "record has: id, domain, project_title, summary, certification, year_completed, "
        "contract_value, duration_months, client_type."
    ))
    doc.add_heading("Domains (5-10 projects each)", level=2)
    bullets(doc, [
        "Engineering & infrastructure: Road Construction, Bridge Engineering, Network Design, Solar Energy, Fleet Management, Medical Equipment",
        "IT & software: Cybersecurity, Hospital IT, ERP Implementation, LMS Development, Custom Software Development, Mobile App Development, Cloud Migration, DevOps & Automation, Data Analytics & BI, AI & Machine Learning, E-Commerce Platforms, Fintech Solutions, SaaS Product Engineering, IT Managed Services",
    ])
    doc.add_heading("Hybrid retrieval", level=2)
    para(doc, (
        "Top-10 candidates are retrieved by dense cosine similarity (fastembed "
        "all-MiniLM-L6-v2 vectors in a FAISS inner-product index), then re-ranked by "
        "hybrid score = 0.5 × dense + 0.5 × structured, where the structured score "
        "weighs domain match (45%), certification match (30%), recency (15%), and "
        "client-type match (10%)."
    ))
    doc.add_page_break()

    # ── 9. LLM strategy ─────────────────────────────────────────────────────
    doc.add_heading("9. LLM Strategy — Caching, Fallback & Demo Mode", level=1)
    bullets(doc, [
        "Cache-first: every LLM call checks a file cache (app/data/cache) keyed by sha256(model + normalized prompt). Hits return instantly; streamed sections replay token-by-token so the typing effect is preserved.",
        "Provider rotation: when GROQ fails (e.g. the 100k tokens/day cap), the call automatically rotates through configured free OpenRouter models until one answers. The response is cached as usual.",
        "DEMO_MODE=true: if every provider fails, the nearest cached response with the same call-site tag is served (similarity-gated) — degraded but never empty.",
        "Deterministic fallbacks: export narrative sections (understanding, case studies, executive summary) degrade to honest, fact-based text rather than failing the export.",
        "scripts/warm_cache.py pre-populates the cache for a known demo document.",
    ])
    doc.add_page_break()

    # ── 10. Schema ──────────────────────────────────────────────────────────
    doc.add_heading("10. Database Schema & Migrations", level=1)
    para(doc, "Six tables in Supabase (full DDL in schema.sql). Children cascade on workspace delete.")
    table(doc, ["Table", "Key columns"], [
        ["workspaces", "id, name, status, rfp_file_path, sector, issuer_type, competitor_presence, pipeline_timings (JSONB), created_at"],
        ["requirements", "id, workspace_id, section_ref, requirement, type, deadline, budget_ref, extraction_source, taxonomy_id/name, note, is_done"],
        ["compliance_items", "id, requirement_id, workspace_id, status (pass/partial/fail), matched_capability_id, confidence, gap_note"],
        ["proposal_sections", "id, workspace_id, requirement_id, section_title, section_ref, ai_draft, edited_draft, status (draft/edited/approved)"],
        ["bid_scores", "id, workspace_id, six factor columns, overall_score, decision, score_breakdown (JSONB)"],
    ])
    doc.add_heading("Pending migration (run once in the Supabase SQL editor)", level=2)
    code(doc, (
        "ALTER TABLE requirements ADD COLUMN IF NOT EXISTS note TEXT;\n"
        "ALTER TABLE requirements ADD COLUMN IF NOT EXISTS is_done BOOLEAN DEFAULT FALSE;"
    ))
    para(doc, (
        "Until these run, saving a requirement note returns an error message containing "
        "exactly this SQL. All other migration notes are kept inline in schema.sql."
    ))
    doc.add_page_break()

    # ── 11. Troubleshooting ─────────────────────────────────────────────────
    doc.add_heading("11. Troubleshooting & Known Behaviours", level=1)
    table(doc, ["Symptom", "Cause & fix"], [
        ["Export button does nothing",
         "Check the toast — it now shows the real server error. If the API is unreachable, "
         "confirm the backend is running on port 8000 and that no stale Python process is "
         "holding the port (Get-NetTCPConnection -LocalPort 8000 -State Listen)."],
        ["GROQ rate-limit errors in logs",
         "The free tier caps at 100k tokens/day. Calls rotate to OpenRouter automatically; "
         "cached responses keep working. Consider a second GROQ key for demos."],
        ["Approve/edit seems to fail on old workspaces",
         "Sections saved before the proposal_sections migration have NULL requirement_id / "
         "section_title. Saving works; regenerate the draft to restore exact per-requirement "
         "evidence citations in the full export (legacy sections fall back to section-group citations)."],
        ["Saving a requirement note returns a database error",
         "Run the two ALTER TABLE statements from section 10 in the Supabase SQL editor."],
        ["Full export is slow the first time",
         "It makes several LLM calls (summary, understanding, 3 case studies). They are "
         "cached — repeat exports of the same workspace take a few seconds."],
        ["Win score looks wrong after manual edits",
         "Re-run POST /score (the competitor-presence dropdown does this automatically)."],
    ])

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
