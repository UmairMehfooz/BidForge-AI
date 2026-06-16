"""Generate a sample Telecom-sector RFP engineered to score a GO decision.

Every mandatory requirement maps onto strong Network Design / Cybersecurity
capability-library evidence; budget (PKR 180M) sits near the average contract
value of the capabilities that should match; sector keywords steer inference
to Telecom (highest historical win rate, 0.65).
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = "Sample_RFP_GO_Broadband_Network.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


# ---- Cover ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("REQUEST FOR PROPOSAL (RFP)\n")
run.bold = True
run.font.size = Pt(20)
run2 = title.add_run("Provincial Broadband Backbone Expansion Programme —\nFiber Optic Network Design, Deployment and Managed NOC Services")
run2.font.size = Pt(14)

p("")
p("Issuing Authority: Provincial Telecommunication & Connectivity Authority (PTCA), Government of Punjab", bold=True)
p("RFP Reference No: PTCA/TELECOM/2026/RFP-014")
p("Date of Issue: 10 June 2026")
p("Sector: Telecommunications / Broadband Connectivity")

h("1. Introduction and Background")
p(
    "The Provincial Telecommunication & Connectivity Authority (PTCA), a government body responsible for "
    "telecommunication policy and public broadband connectivity, invites sealed proposals from reputable firms "
    "for the design, deployment, and managed operation of a provincial fiber optic broadband backbone. The "
    "programme will extend high-capacity fiber connectivity to government facilities, public institutions, and "
    "underserved urban centers, and will establish a centrally managed Network Operations Center (NOC) to "
    "guarantee telecommunication service quality across the province."
)
p(
    "This telecommunication initiative forms part of the National Broadband Strategy and will interconnect with "
    "existing telecom carrier infrastructure. Bidders are expected to demonstrate proven network design "
    "experience on large government fiber and broadband projects of comparable scale."
)

h("2. Scope of Work")
bullet("Detailed network design for approximately 150 km of fiber optic backbone interconnecting 12 government facilities and 3 telecom exchange points, with fully redundant ring topology.")
bullet("Supply, installation, and commissioning of optical line terminal equipment, core and edge routing, and broadband aggregation nodes.")
bullet("Design and deployment of a metropolitan public Wi-Fi access network in two district headquarters, integrated with the fiber backbone.")
bullet("Establishment of a 24/7 Network Operations Center (NOC) with proactive monitoring, fault management, and a guaranteed 99.9% service-level agreement (SLA).")
bullet("Implementation of a cybersecurity framework for the backbone, including next-generation firewalls and SIEM-based security monitoring.")
bullet("Structured training and certification of PTCA network administrators on the deployed network design and NOC tooling.")

h("3. Mandatory Eligibility and Technical Requirements")
p("Bidders must comply with ALL of the following mandatory requirements. Non-compliance with any mandatory clause will result in disqualification.", bold=True)
p("3.1 The bidder must hold a valid ISO 27001 certification for information security management.")
p("3.2 The bidder must hold a valid CMMI Level 3 (or higher) appraisal for process maturity.")
p("3.3 The bidder must have successfully completed at least one fiber optic backbone network design and deployment project of 100 km or more for a federal or provincial government client within the last five years.")
p("3.4 The bidder must demonstrate experience establishing a Network Operations Center (NOC) with a committed SLA of 99.9% or higher on a telecommunication or broadband network.")
p("3.5 The bidder must have completed at least one network design contract with a value of PKR 150 Million or above.")
p("3.6 The bidder must demonstrate experience deploying a large-scale public Wi-Fi or metropolitan wireless mesh network serving at least 100,000 users.")
p("3.7 The bidder must demonstrate experience implementing network redundancy and DDoS protection on a carrier-grade or ISP backbone network.")
p("3.8 The bidder must demonstrate experience delivering structured training programmes for client network administrators as part of a network deployment engagement.")
p("3.9 The bidder must demonstrate cybersecurity implementation experience, including firewall deployment and SIEM-based security operations, on government or enterprise networks.")
p("3.10 The bidder must provide network design documentation deliverables (high-level design, low-level design, and as-built drawings) as part of the engagement.")

h("4. Questions to Bidders")
p("Q1: Describe your proposed network design methodology for a redundant fiber optic backbone, including how you will guarantee the 99.9% SLA for telecommunication services.")
p("Q2: Describe the staffing model, monitoring toolchain, and escalation procedures for the proposed 24/7 Network Operations Center.")
p("Q3: Describe a comparable government broadband or fiber backbone project you have delivered, including scale, outcomes, and lessons learned.")

h("5. Budget")
p(
    "The total estimated budget for this telecommunication programme is PKR 180 Million, inclusive of all "
    "design, equipment, deployment, NOC establishment, and first-year managed services. Financial proposals "
    "exceeding the estimated budget may be rejected."
)

h("6. Evaluation Criteria")
p("Proposals will be evaluated against the following weighted criteria:")
bullet("Technical Approach and Network Design Quality — 40%")
bullet("Relevant Experience on Comparable Telecom/Broadband Projects — 30%")
bullet("Financial Proposal — 20%")
bullet("Certifications and Quality Assurance (ISO 27001, CMMI L3) — 10%")

h("7. Submission Instructions")
p(
    "Sealed technical and financial proposals must be submitted to the PTCA Procurement Wing, Lahore, "
    "no later than 20 July 2026 at 1500 hours PKT. Late submissions will not be accepted. Bid validity "
    "shall be 120 days from the submission deadline."
)
p("A pre-bid telecom industry briefing will be held on 25 June 2026 at the PTCA head office.")

doc.save(OUT)
print(f"Saved {OUT}")
