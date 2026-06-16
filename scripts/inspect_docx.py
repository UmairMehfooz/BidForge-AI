"""Dump heading structure + flagged gaps of a DOCX (debug helper)."""
import sys
from docx import Document

doc = Document(sys.argv[1])
gap_count = 0
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        level = p.style.name.replace("Heading ", "H")
        print(f"{level}: {p.text[:90]}")
    elif "[GAP FLAGGED" in p.text:
        gap_count += 1
print(f"\nTables: {len(doc.tables)}")
if doc.tables:
    t = doc.tables[-1]
    print(f"Last table: {len(t.rows)} rows x {len(t.columns)} cols")
    print("Header:", " | ".join(c.text for c in t.rows[0].cells))
print(f"Gap flags: {gap_count}")
